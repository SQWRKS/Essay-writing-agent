import json
import os
from app.models import Project


def _coerce_section_content(section_payload):
    if isinstance(section_payload, dict):
        return str(section_payload.get("content") or "")
    if isinstance(section_payload, str):
        return section_payload
    return str(section_payload or "")


def export_project_docx(project: Project, output_dir: str) -> str:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    content = {}
    if project.content:
        try:
            content = json.loads(project.content)
        except Exception:
            content = {}

    doc = Document()

    title = doc.add_heading(project.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Topic: {project.topic}")
    doc.add_paragraph(f"Status: {project.status}")
    doc.add_paragraph()

    sections = content.get("sections", {})
    for sec_key, sec_content in sections.items():
        heading = sec_key.replace("_", " ").title()
        doc.add_heading(heading, level=1)
        section_text = _coerce_section_content(sec_content)
        for para in section_text.split("\n\n"):
            cleaned = para.strip()
            if not cleaned:
                continue
            if cleaned.startswith("## "):
                doc.add_heading(cleaned[3:].strip(), level=2)
                continue
            p = doc.add_paragraph(cleaned)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    metadata = content.get("metadata", {})
    figures = metadata.get("figures", [])
    if figures:
        doc.add_heading("Figures", level=1)
        for idx, fig in enumerate(figures, 1):
            title = fig.get("title", f"Figure {idx}")
            doc.add_paragraph(title)

            fig_path = fig.get("path", "")
            if fig_path and os.path.exists(fig_path):
                try:
                    doc.add_picture(fig_path, width=Inches(6.0))
                except Exception:
                    pass
            elif fig.get("url"):
                doc.add_paragraph(f"Reference: {fig['url']}")

            description = fig.get("description", "")
            if description:
                doc.add_paragraph(description)

    bibliography = metadata.get("bibliography", "")
    if bibliography:
        doc.add_heading("Bibliography", level=1)
        for line in bibliography.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    filename = f"project_{project.id[:8]}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath
